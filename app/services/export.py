"""Export service - Generate DOCX and HTML exports of search/compare/QA results."""

import html
import re
from datetime import datetime
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.models import SearchResult, QAResponse, Citation


# =============================================================================
# HTML Export Functions
# =============================================================================

def _get_html_base_styles() -> str:
    """Return base CSS styles for HTML exports."""
    return """
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #1f2937;
            line-height: 1.6;
        }
        h1 {
            color: #111827;
            border-bottom: 3px solid #dc2626;
            padding-bottom: 12px;
            margin-bottom: 24px;
        }
        h2 {
            color: #374151;
            margin-top: 32px;
            margin-bottom: 16px;
        }
        .meta {
            color: #6b7280;
            font-size: 14px;
            margin-bottom: 24px;
        }
        .result-card {
            background: #f9fafb;
            border: 1px solid #e5e7eb;
            border-radius: 8px;
            padding: 20px;
            margin-bottom: 16px;
        }
        .result-header {
            display: flex;
            justify-content: space-between;
            align-items: flex-start;
            margin-bottom: 12px;
        }
        .result-title {
            font-weight: 600;
            color: #111827;
            font-size: 16px;
        }
        .result-page {
            color: #6b7280;
            font-size: 14px;
        }
        .result-score {
            color: #9ca3af;
            font-size: 12px;
        }
        .snippet {
            color: #374151;
            font-size: 14px;
            line-height: 1.7;
        }
        .snippet mark {
            background-color: #fef08a;
            padding: 1px 3px;
            border-radius: 2px;
        }
        .citation-box {
            background: #eff6ff;
            border-left: 4px solid #3b82f6;
            padding: 16px;
            margin: 12px 0;
        }
        .citation-header {
            font-weight: 600;
            color: #1e40af;
            margin-bottom: 8px;
        }
        .citation-text {
            color: #1e3a5f;
            font-style: italic;
            font-size: 14px;
        }
        .answer-box {
            background: #f0fdf4;
            border: 1px solid #bbf7d0;
            border-radius: 8px;
            padding: 20px;
            margin: 20px 0;
        }
        .answer-text {
            white-space: pre-line;
            color: #166534;
        }
        .warning-box {
            background: #fffbeb;
            border-left: 4px solid #f59e0b;
            padding: 16px;
            margin: 16px 0;
        }
        .doc-group {
            margin-bottom: 24px;
        }
        .doc-group-header {
            background: #dbeafe;
            padding: 12px 16px;
            border-radius: 6px 6px 0 0;
            font-weight: 600;
            color: #1e40af;
        }
        .doc-group-content {
            border: 1px solid #e5e7eb;
            border-top: none;
            border-radius: 0 0 6px 6px;
            padding: 16px;
        }
        .match-item {
            border-left: 2px solid #e5e7eb;
            padding-left: 12px;
            margin-bottom: 12px;
        }
        .match-page {
            font-weight: 500;
            color: #374151;
            font-size: 14px;
        }
        .match-snippet {
            color: #6b7280;
            font-size: 14px;
            margin-top: 4px;
        }
        @media print {
            body { padding: 0; }
            .result-card { break-inside: avoid; }
            .citation-box { break-inside: avoid; }
        }
    """


def _html_escape(text: str) -> str:
    """Escape HTML special characters but preserve <mark> tags for highlighting."""
    # First escape all HTML
    escaped = html.escape(text)
    # Then restore <mark> tags (they're safe for highlighting)
    escaped = escaped.replace('&lt;mark&gt;', '<mark>')
    escaped = escaped.replace('&lt;/mark&gt;', '</mark>')
    return escaped


def export_search_results_html(results: list[SearchResult], query: str) -> str:
    """
    Export search results to standalone HTML.

    Args:
        results: List of SearchResult objects
        query: The search query string

    Returns:
        Complete HTML document as string
    """
    timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")

    results_html = ""
    if results:
        for result in results:
            results_html += f"""
            <div class="result-card">
                <div class="result-header">
                    <div>
                        <div class="result-title">{_html_escape(result.filename)}</div>
                        <div class="result-page">Page {result.page_number}</div>
                    </div>
                    <div class="result-score">Score: {result.score:.2f}</div>
                </div>
                <div class="snippet">{_html_escape(result.snippet)}</div>
            </div>
            """
    else:
        results_html = """
            <div class="warning-box">
                No results found for this search query.
            </div>
        """

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Search Results: {_html_escape(query)}</title>
    <style>{_get_html_base_styles()}</style>
</head>
<body>
    <h1>Search Results</h1>
    <div class="meta">
        <strong>Query:</strong> "{_html_escape(query)}"<br>
        <strong>Results:</strong> {len(results)} found<br>
        <strong>Exported:</strong> {timestamp}
    </div>

    {results_html}

    <div class="meta" style="margin-top: 32px; border-top: 1px solid #e5e7eb; padding-top: 16px;">
        Generated by Contract Dashboard
    </div>
</body>
</html>"""


def export_compare_results_html(comparison: dict, topic: str) -> str:
    """
    Export comparison results to standalone HTML.

    Args:
        comparison: Comparison result dict from compare_documents or compare_documents_multi
        topic: The search topic/term

    Returns:
        Complete HTML document as string
    """
    timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")

    content_html = ""

    # Handle multi-document comparison format
    if "documents" in comparison and "matches" in comparison:
        documents = comparison.get("documents", [])
        matches = comparison.get("matches", [])

        # Document list
        doc_list = ", ".join([doc.get("filename", "Unknown") for doc in documents])

        content_html += f"""
        <h2>Documents Compared</h2>
        <p class="meta">{doc_list}</p>
        """

        if matches:
            # Group matches by filename
            matches_by_file: dict[str, list] = {}
            for match in matches:
                filename = match.get("filename", "Unknown")
                if filename not in matches_by_file:
                    matches_by_file[filename] = []
                matches_by_file[filename].append(match)

            content_html += "<h2>Matches</h2>"

            for filename, file_matches in matches_by_file.items():
                content_html += f"""
                <div class="doc-group">
                    <div class="doc-group-header">{_html_escape(filename)}</div>
                    <div class="doc-group-content">
                """
                for match in file_matches:
                    snippet = match.get("snippet", "")
                    page_num = match.get("page_number", match.get("page", "?"))
                    content_html += f"""
                        <div class="match-item">
                            <div class="match-page">Page {page_num}</div>
                            <div class="match-snippet">{_html_escape(snippet)}</div>
                        </div>
                    """
                content_html += "</div></div>"
        else:
            content_html += """
            <div class="warning-box">
                No matches found for the specified topic in the selected documents.
            </div>
            """

    # Handle legacy two-document comparison format
    elif "doc_a" in comparison and "doc_b" in comparison:
        doc_a = comparison.get("doc_a", {})
        doc_b = comparison.get("doc_b", {})
        matches_a = comparison.get("matches_a", [])
        matches_b = comparison.get("matches_b", [])

        for doc, matches, label in [
            (doc_a, matches_a, "Document A"),
            (doc_b, matches_b, "Document B"),
        ]:
            filename = doc.get("filename", "Unknown")
            content_html += f"""
            <div class="doc-group">
                <div class="doc-group-header">{label}: {_html_escape(filename)}</div>
                <div class="doc-group-content">
            """
            if matches:
                for match in matches:
                    page_num = match.get("page", "?")
                    context = match.get("context", "")
                    content_html += f"""
                        <div class="match-item">
                            <div class="match-page">Page {page_num}</div>
                            <div class="match-snippet">{_html_escape(context)}</div>
                        </div>
                    """
            else:
                content_html += '<p class="meta">No matches found in this document.</p>'
            content_html += "</div></div>"

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document Comparison: {_html_escape(topic) if topic else 'All Content'}</title>
    <style>{_get_html_base_styles()}</style>
</head>
<body>
    <h1>Document Comparison</h1>
    <div class="meta">
        <strong>Topic:</strong> "{_html_escape(topic) if topic else 'N/A'}"<br>
        <strong>Exported:</strong> {timestamp}
    </div>

    {content_html}

    <div class="meta" style="margin-top: 32px; border-top: 1px solid #e5e7eb; padding-top: 16px;">
        Generated by Contract Dashboard
    </div>
</body>
</html>"""


def export_qa_response_html(response: QAResponse, question: str) -> str:
    """
    Export Q&A response to standalone HTML.

    Args:
        response: QAResponse object with answer and citations
        question: The original question

    Returns:
        Complete HTML document as string
    """
    timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")

    answer_html = ""
    if response.no_evidence:
        answer_html = f"""
        <div class="warning-box">
            <strong>Not found in documents</strong><br>
            {_html_escape(response.answer)}
        </div>
        """
    else:
        answer_html = f"""
        <div class="answer-box">
            <div class="answer-text">{_html_escape(response.answer)}</div>
        </div>
        """

    citations_html = ""
    if response.citations:
        citations_html = "<h2>Sources</h2>"
        for i, citation in enumerate(response.citations, 1):
            citations_html += f"""
            <div class="citation-box">
                <div class="citation-header">
                    [{i}] {_html_escape(citation.filename)} - Page {citation.page_number}
                </div>
                <div class="citation-text">
                    "{_html_escape(citation.cited_text)}{'...' if len(citation.cited_text) >= 200 else ''}"
                </div>
            </div>
            """

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Q&A: {_html_escape(question[:50])}{'...' if len(question) > 50 else ''}</title>
    <style>{_get_html_base_styles()}</style>
</head>
<body>
    <h1>Q&A Response</h1>
    <div class="meta">
        <strong>Question:</strong> "{_html_escape(question)}"<br>
        <strong>Exported:</strong> {timestamp}
    </div>

    <h2>Answer</h2>
    {answer_html}

    {citations_html}

    <div class="meta" style="margin-top: 32px; border-top: 1px solid #e5e7eb; padding-top: 16px;">
        Generated by Contract Dashboard
    </div>
</body>
</html>"""


# =============================================================================
# DOCX Export Functions
# =============================================================================

def _strip_html_tags(text: str) -> str:
    """Remove HTML tags from text (used for DOCX which doesn't support HTML)."""
    # Replace <mark> tags with nothing (we'll handle highlighting differently)
    text = re.sub(r'<mark>', '', text)
    text = re.sub(r'</mark>', '', text)
    # Remove any other HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    return html.unescape(text)


def _setup_docx_styles(doc: Document) -> None:
    """Set up custom styles for the document."""
    styles = doc.styles

    # Title style
    if 'CustomTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(17, 24, 39)
        title_style.paragraph_format.space_after = Pt(12)

    # Heading style
    if 'CustomHeading' not in [s.name for s in styles]:
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(55, 65, 81)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)

    # Meta info style
    if 'MetaInfo' not in [s.name for s in styles]:
        meta_style = styles.add_style('MetaInfo', WD_STYLE_TYPE.PARAGRAPH)
        meta_style.font.size = Pt(10)
        meta_style.font.color.rgb = RGBColor(107, 114, 128)
        meta_style.paragraph_format.space_after = Pt(4)

    # Citation style
    if 'CitationStyle' not in [s.name for s in styles]:
        citation_style = styles.add_style('CitationStyle', WD_STYLE_TYPE.PARAGRAPH)
        citation_style.font.size = Pt(10)
        citation_style.font.italic = True
        citation_style.font.color.rgb = RGBColor(30, 58, 95)


def _add_header_footer(doc: Document, title: str) -> None:
    """Add header and footer to document."""
    section = doc.sections[0]

    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"Contract Dashboard - {title}"
    header_para.style.font.size = Pt(9)
    header_para.style.font.color.rgb = RGBColor(156, 163, 175)

    # Footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.style.font.size = Pt(9)


def export_search_results_docx(results: list[SearchResult], query: str) -> bytes:
    """
    Export search results to DOCX format.

    Args:
        results: List of SearchResult objects
        query: The search query string

    Returns:
        DOCX document as bytes
    """
    doc = Document()
    _setup_docx_styles(doc)
    _add_header_footer(doc, "Search Results")

    timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")

    # Title
    title = doc.add_paragraph("Search Results", style='CustomTitle')

    # Meta information
    doc.add_paragraph(f'Query: "{query}"', style='MetaInfo')
    doc.add_paragraph(f"Results: {len(results)} found", style='MetaInfo')
    doc.add_paragraph(f"Exported: {timestamp}", style='MetaInfo')

    # Add horizontal line
    doc.add_paragraph("_" * 80)

    if results:
        for i, result in enumerate(results, 1):
            # Result header
            header = doc.add_paragraph(style='CustomHeading')
            header.add_run(f"{i}. {result.filename}").bold = True
            header.add_run(f" - Page {result.page_number}")

            # Score
            score_para = doc.add_paragraph(style='MetaInfo')
            score_para.add_run(f"Relevance Score: {result.score:.2f}")

            # Snippet (strip HTML tags)
            snippet_text = _strip_html_tags(result.snippet)
            snippet_para = doc.add_paragraph()
            snippet_para.add_run(snippet_text)
            snippet_para.paragraph_format.left_indent = Inches(0.25)

            # Add spacing
            doc.add_paragraph()
    else:
        warning = doc.add_paragraph()
        warning.add_run("No results found for this search query.").italic = True

    # Footer note
    doc.add_paragraph("_" * 80)
    footer_note = doc.add_paragraph(style='MetaInfo')
    footer_note.add_run("Generated by Contract Dashboard")
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_compare_results_docx(comparison: dict, topic: str) -> bytes:
    """
    Export comparison results to DOCX format.

    Args:
        comparison: Comparison result dict
        topic: The search topic/term

    Returns:
        DOCX document as bytes
    """
    doc = Document()
    _setup_docx_styles(doc)
    _add_header_footer(doc, "Document Comparison")

    timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")

    # Title
    doc.add_paragraph("Document Comparison", style='CustomTitle')

    # Meta information
    doc.add_paragraph(f'Topic: "{topic if topic else "N/A"}"', style='MetaInfo')
    doc.add_paragraph(f"Exported: {timestamp}", style='MetaInfo')

    doc.add_paragraph("_" * 80)

    # Handle multi-document comparison format
    if "documents" in comparison and "matches" in comparison:
        documents = comparison.get("documents", [])
        matches = comparison.get("matches", [])

        # Documents compared
        doc.add_paragraph("Documents Compared", style='CustomHeading')
        for document in documents:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.add_run(document.get("filename", "Unknown"))

        if matches:
            doc.add_paragraph("Matches by Document", style='CustomHeading')

            # Group matches by filename
            matches_by_file: dict[str, list] = {}
            for match in matches:
                filename = match.get("filename", "Unknown")
                if filename not in matches_by_file:
                    matches_by_file[filename] = []
                matches_by_file[filename].append(match)

            for filename, file_matches in matches_by_file.items():
                # Document subheading
                doc_header = doc.add_paragraph()
                doc_header.add_run(filename).bold = True

                for match in file_matches:
                    page_num = match.get("page_number", match.get("page", "?"))
                    snippet = _strip_html_tags(match.get("snippet", ""))

                    match_para = doc.add_paragraph()
                    match_para.add_run(f"Page {page_num}: ").bold = True
                    match_para.add_run(snippet)
                    match_para.paragraph_format.left_indent = Inches(0.25)

                doc.add_paragraph()
        else:
            warning = doc.add_paragraph()
            warning.add_run("No matches found for the specified topic.").italic = True

    # Handle legacy two-document comparison format
    elif "doc_a" in comparison and "doc_b" in comparison:
        doc_a = comparison.get("doc_a", {})
        doc_b = comparison.get("doc_b", {})
        matches_a = comparison.get("matches_a", [])
        matches_b = comparison.get("matches_b", [])

        for document, matches, label in [
            (doc_a, matches_a, "Document A"),
            (doc_b, matches_b, "Document B"),
        ]:
            filename = document.get("filename", "Unknown")

            doc.add_paragraph(f"{label}: {filename}", style='CustomHeading')

            if matches:
                for match in matches:
                    page_num = match.get("page", "?")
                    context = _strip_html_tags(match.get("context", ""))

                    match_para = doc.add_paragraph()
                    match_para.add_run(f"Page {page_num}: ").bold = True
                    match_para.add_run(context)
                    match_para.paragraph_format.left_indent = Inches(0.25)
            else:
                no_match = doc.add_paragraph()
                no_match.add_run("No matches found in this document.").italic = True

            doc.add_paragraph()

    # Footer note
    doc.add_paragraph("_" * 80)
    footer_note = doc.add_paragraph(style='MetaInfo')
    footer_note.add_run("Generated by Contract Dashboard")
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_qa_response_docx(response: QAResponse, question: str) -> bytes:
    """
    Export Q&A response to DOCX format.

    Args:
        response: QAResponse object with answer and citations
        question: The original question

    Returns:
        DOCX document as bytes
    """
    doc = Document()
    _setup_docx_styles(doc)
    _add_header_footer(doc, "Q&A Response")

    timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")

    # Title
    doc.add_paragraph("Q&A Response", style='CustomTitle')

    # Meta information
    doc.add_paragraph(f'Question: "{question}"', style='MetaInfo')
    doc.add_paragraph(f"Exported: {timestamp}", style='MetaInfo')

    doc.add_paragraph("_" * 80)

    # Answer section
    doc.add_paragraph("Answer", style='CustomHeading')

    if response.no_evidence:
        warning = doc.add_paragraph()
        warning.add_run("Not found in documents: ").bold = True
        warning.add_run(response.answer)
    else:
        answer_para = doc.add_paragraph()
        answer_para.add_run(response.answer)
        answer_para.paragraph_format.left_indent = Inches(0.25)

    # Citations section
    if response.citations:
        doc.add_paragraph("Sources", style='CustomHeading')

        for i, citation in enumerate(response.citations, 1):
            # Citation header
            cite_header = doc.add_paragraph()
            cite_header.add_run(f"[{i}] {citation.filename}").bold = True
            cite_header.add_run(f" - Page {citation.page_number}")

            # Cited text
            cite_text = doc.add_paragraph(style='CitationStyle')
            text = citation.cited_text
            if len(text) >= 200:
                text += "..."
            cite_text.add_run(f'"{text}"')
            cite_text.paragraph_format.left_indent = Inches(0.25)

            doc.add_paragraph()

    # Footer note
    doc.add_paragraph("_" * 80)
    footer_note = doc.add_paragraph(style='MetaInfo')
    footer_note.add_run("Generated by Contract Dashboard")
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
