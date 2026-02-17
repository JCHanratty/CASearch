"""Compare routes - side-by-side document comparison."""

from typing import Optional
from fastapi import APIRouter, Request, Query
from fastapi.responses import HTMLResponse, Response

from app.db import get_db
from app.services.compare import compare_documents, compare_documents_multi
from app.services.compare_ai import ai_compare_documents
from app.services.export import export_compare_results_html, export_compare_results_docx
from app.templates import templates

router = APIRouter()


@router.get("/", response_class=HTMLResponse)
async def compare_page(request: Request):
    """Compare documents page."""
    # Get list of indexed documents for selection
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id, filename FROM files WHERE status = 'indexed' ORDER BY filename"
        ).fetchall()
        documents = [{"id": r["id"], "filename": r["filename"]} for r in rows]

    return templates.TemplateResponse(
        "compare.html",
        {
            "request": request,
            "documents": documents,
            "page_title": "Compare",
        },
    )


@router.get("/results", response_class=HTMLResponse)
async def compare_results(
    request: Request,
    doc_a: int = 0,
    doc_b: int = 0,
    topic: str = "",
    doc_ids: Optional[list[int]] = Query(None),
    mode: str = "text",  # "text" or "ai"
):
    """Get comparison results."""

    result = None
    multi_result = None
    ai_result = None

    # AI comparison mode
    if mode == "ai":
        if doc_ids and len(doc_ids) >= 2:
            ai_result = ai_compare_documents(doc_ids, topic if topic.strip() else None)
        elif doc_a and doc_b and doc_a != doc_b:
            ai_result = ai_compare_documents([doc_a, doc_b], topic if topic.strip() else None)
    else:
        # Text comparison mode (default)
        # Multi-select mode: doc_ids takes precedence
        if doc_ids and len(doc_ids) >= 2:
            multi_result = compare_documents_multi(doc_ids, topic if topic.strip() else None)
        # Legacy two-document mode
        elif doc_a and doc_b and doc_a != doc_b:
            result = compare_documents(doc_a, doc_b, topic if topic.strip() else None)

    # Check if HTMX request (partial) or full page
    if request.headers.get("HX-Request"):
        if mode == "ai":
            return templates.TemplateResponse(
                "components/compare_ai_results.html",
                {
                    "request": request,
                    "ai_result": ai_result,
                    "topic": topic,
                },
            )
        else:
            return templates.TemplateResponse(
                "components/compare_results.html",
                {
                    "request": request,
                    "result": result,
                    "multi_result": multi_result,
                    "topic": topic,
                },
            )

    # Get list of indexed documents for selection
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id, filename FROM files WHERE status = 'indexed' ORDER BY filename"
        ).fetchall()
        documents = [{"id": r["id"], "filename": r["filename"]} for r in rows]

    return templates.TemplateResponse(
        "compare.html",
        {
            "request": request,
            "documents": documents,
            "result": result,
            "multi_result": multi_result,
            "ai_result": ai_result,
            "topic": topic,
            "mode": mode,
            "page_title": "Compare",
        },
    )


@router.get("/export")
async def export_compare(
    doc_a: int = 0,
    doc_b: int = 0,
    topic: str = "",
    doc_ids: Optional[list[int]] = Query(None),
    format: str = Query("html", regex="^(html|docx)$"),
):
    """
    Export comparison results to HTML or DOCX format.

    Args:
        doc_a: First document ID (legacy mode)
        doc_b: Second document ID (legacy mode)
        topic: Search topic/term
        doc_ids: List of document IDs (multi-select mode)
        format: Export format ("html" or "docx")

    Returns:
        HTML or DOCX file download
    """
    comparison = {}

    # Multi-select mode: doc_ids takes precedence
    if doc_ids and len(doc_ids) >= 2:
        comparison = compare_documents_multi(doc_ids, topic if topic.strip() else None)
    # Legacy two-document mode
    elif doc_a and doc_b and doc_a != doc_b:
        result = compare_documents(doc_a, doc_b, topic if topic.strip() else None)
        if result:
            comparison = result

    topic_slug = topic[:20].replace(' ', '_') if topic else "comparison"

    if format == "docx":
        content = export_compare_results_docx(comparison, topic)
        filename = f"compare_{topic_slug}.docx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    else:
        content = export_compare_results_html(comparison, topic)
        filename = f"compare_{topic_slug}.html"
        return Response(
            content=content,
            media_type="text/html",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )


@router.get("/export-ai")
async def export_ai_compare(
    doc_a: int = 0,
    doc_b: int = 0,
    topic: str = "",
    doc_ids: Optional[list[int]] = Query(None),
    format: str = Query("html", regex="^(html|docx)$"),
):
    """
    Export AI comparison results to HTML or DOCX format.

    Args:
        doc_a: First document ID (legacy mode)
        doc_b: Second document ID (legacy mode)
        topic: Search topic/term
        doc_ids: List of document IDs (multi-select mode)
        format: Export format ("html" or "docx")

    Returns:
        HTML or DOCX file download
    """
    ai_result = None

    # Get AI comparison result
    if doc_ids and len(doc_ids) >= 2:
        ai_result = ai_compare_documents(doc_ids, topic if topic.strip() else None)
    elif doc_a and doc_b and doc_a != doc_b:
        ai_result = ai_compare_documents([doc_a, doc_b], topic if topic.strip() else None)

    topic_slug = topic[:20].replace(' ', '_') if topic else "ai_comparison"

    if format == "docx":
        content = _export_ai_results_docx(ai_result, topic)
        filename = f"ai_compare_{topic_slug}.docx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    else:
        content = _export_ai_results_html(ai_result, topic)
        filename = f"ai_compare_{topic_slug}.html"
        return Response(
            content=content,
            media_type="text/html",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )


def _export_ai_results_html(ai_result: dict, topic: str) -> str:
    """Generate HTML export for AI comparison results."""
    if not ai_result:
        return "<html><body><p>No AI comparison results available.</p></body></html>"

    analysis = ai_result.get("analysis", "No analysis available.")
    sources = ai_result.get("sources", [])

    sources_html = ""
    if sources:
        sources_html = "<h2>Sources</h2><ul>"
        for source in sources:
            sources_html += f"<li>{source.get('filename', 'Unknown')} (ID: {source.get('id', 'N/A')})</li>"
        sources_html += "</ul>"

    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>AI Comparison: {topic or 'Document Analysis'}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1 {{ color: #333; border-bottom: 2px solid #007bff; padding-bottom: 10px; }}
        h2 {{ color: #555; margin-top: 30px; }}
        .analysis {{ background: #f8f9fa; padding: 20px; border-radius: 8px; white-space: pre-wrap; }}
        ul {{ padding-left: 20px; }}
        li {{ margin: 5px 0; }}
    </style>
</head>
<body>
    <h1>AI Document Comparison{f': {topic}' if topic else ''}</h1>
    <h2>Analysis</h2>
    <div class="analysis">{analysis}</div>
    {sources_html}
</body>
</html>"""


def _export_ai_results_docx(ai_result: dict, topic: str) -> bytes:
    """Generate DOCX export for AI comparison results."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(f"AI Document Comparison{f': {topic}' if topic else ''}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not ai_result:
        doc.add_paragraph("No AI comparison results available.")
    else:
        # Analysis section
        doc.add_heading("Analysis", level=1)
        analysis = ai_result.get("analysis", "No analysis available.")
        doc.add_paragraph(analysis)

        # Sources section
        sources = ai_result.get("sources", [])
        if sources:
            doc.add_heading("Sources", level=1)
            for source in sources:
                doc.add_paragraph(
                    f"• {source.get('filename', 'Unknown')} (ID: {source.get('id', 'N/A')})",
                    style='List Bullet'
                )

    # Save to bytes
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
