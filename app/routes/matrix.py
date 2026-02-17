"""Matrix Compare routes - multi-document aspect-based comparison grid."""

import csv
import html
import re
from io import BytesIO, StringIO

from fastapi import APIRouter, Form, Request
from fastapi.responses import HTMLResponse, Response

from app.services.compare_matrix import compare_matrix
from app.services.file_scanner import get_all_files
from app.templates import templates

router = APIRouter()


@router.get("/", response_class=HTMLResponse)
async def matrix_page(request: Request):
    """Render the Matrix Compare page with all indexed files listed."""
    files = get_all_files()
    indexed_files = [f for f in files if f.status == "indexed"]

    return templates.TemplateResponse(
        "matrix.html",
        {
            "request": request,
            "documents": indexed_files,
            "page_title": "Matrix Compare",
        },
    )


@router.post("/compare", response_class=HTMLResponse)
async def matrix_compare(
    request: Request,
    topic: str = Form(""),
    file_ids: list[int] = Form([]),
):
    """
    Run a matrix comparison across selected documents for a given topic.

    Accepts form data with topic and file_ids, calls compare_matrix(),
    and returns the matrix_results.html partial for HTMX swap.
    """
    error = None
    matrix_result = None

    if not topic.strip():
        error = "Please enter a topic to compare."
    elif len(file_ids) < 2:
        error = "Please select at least 2 documents to compare."
    else:
        try:
            matrix_result = compare_matrix(topic.strip(), file_ids)
            if matrix_result.get("error"):
                error = matrix_result["error"]
                matrix_result = None
        except Exception as e:
            error = f"Matrix comparison failed: {e}"

    return templates.TemplateResponse(
        "components/matrix_results.html",
        {
            "request": request,
            "matrix": matrix_result,
            "topic": topic.strip(),
            "file_ids": file_ids,
            "error": error,
        },
    )


@router.post("/export/{fmt}", response_class=Response)
async def matrix_export(
    fmt: str,
    request: Request,
    topic: str = Form(""),
    file_ids: list[int] = Form([]),
):
    """
    Export matrix comparison results to CSV, HTML, or DOCX.

    Re-runs the comparison (will hit cache) and converts the result into
    the requested download format.
    """
    if fmt not in ("csv", "html", "docx"):
        return Response(content="Unsupported format.", status_code=400)

    if not topic.strip() or len(file_ids) < 2:
        return Response(
            content="A topic and at least 2 documents are required.", status_code=400
        )

    try:
        matrix_result = compare_matrix(topic.strip(), file_ids)
        if matrix_result.get("error"):
            return Response(content=matrix_result["error"], status_code=500)
    except Exception as e:
        return Response(content=f"Comparison failed: {e}", status_code=500)

    topic_slug = re.sub(r"[^a-zA-Z0-9]+", "_", topic.strip()[:30]).strip("_") or "matrix"

    # ----- CSV -----
    if fmt == "csv":
        content = _export_csv(matrix_result)
        return Response(
            content=content,
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="matrix_{topic_slug}.csv"'},
        )

    # ----- HTML -----
    if fmt == "html":
        content = _export_html(matrix_result)
        return Response(
            content=content,
            media_type="text/html",
            headers={"Content-Disposition": f'attachment; filename="matrix_{topic_slug}.html"'},
        )

    # ----- DOCX -----
    if fmt == "docx":
        content = _export_docx(matrix_result)
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="matrix_{topic_slug}.docx"'},
        )


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------

def _export_csv(matrix_result: dict) -> str:
    """Generate a CSV string from the matrix result."""
    aspects = matrix_result.get("aspects", [])
    documents = matrix_result.get("documents", [])
    topic = matrix_result.get("topic", "")

    buf = StringIO()
    writer = csv.writer(buf)

    # Header row: Aspect | Doc1 | Doc2 | ...
    doc_names = [d["name"] for d in documents]
    writer.writerow(["Aspect"] + doc_names)

    # Data rows
    for aspect in aspects:
        row = [aspect]
        for doc in documents:
            row.append(doc.get("values", {}).get(aspect, ""))
        writer.writerow(row)

    return buf.getvalue()


def _export_html(matrix_result: dict) -> str:
    """Generate a standalone HTML document from the matrix result."""
    aspects = matrix_result.get("aspects", [])
    documents = matrix_result.get("documents", [])
    topic = matrix_result.get("topic", "")
    doc_names = [d["name"] for d in documents]

    header_cells = "".join(f"<th>{html.escape(n)}</th>" for n in doc_names)
    rows_html = ""
    for aspect in aspects:
        cells = ""
        for doc in documents:
            value = doc.get("values", {}).get(aspect, "")
            cells += f"<td>{html.escape(value)}</td>"
        rows_html += f"<tr><td class='aspect'>{html.escape(aspect)}</td>{cells}</tr>"

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Matrix Compare: {html.escape(topic)}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
               max-width: 1100px; margin: 0 auto; padding: 40px 20px; color: #1f2937; }}
        h1 {{ color: #111827; border-bottom: 3px solid #dc2626; padding-bottom: 12px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 24px; }}
        th, td {{ border: 1px solid #e5e7eb; padding: 10px 14px; text-align: left; font-size: 14px; }}
        th {{ background: #f1f5f9; color: #334155; font-weight: 600; }}
        td.aspect {{ font-weight: 600; background: #f8fafc; white-space: nowrap; }}
        tr:nth-child(even) {{ background: #f9fafb; }}
        .meta {{ color: #6b7280; font-size: 13px; margin-bottom: 16px; }}
    </style>
</head>
<body>
    <h1>Matrix Compare: {html.escape(topic)}</h1>
    <p class="meta">Documents: {', '.join(html.escape(n) for n in doc_names)}</p>
    <table>
        <thead><tr><th>Aspect</th>{header_cells}</tr></thead>
        <tbody>{rows_html}</tbody>
    </table>
    <p class="meta" style="margin-top:32px;border-top:1px solid #e5e7eb;padding-top:16px;">
        Generated by Contract Dashboard
    </p>
</body>
</html>"""


def _export_docx(matrix_result: dict) -> bytes:
    """Generate a DOCX file with the matrix table."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    aspects = matrix_result.get("aspects", [])
    documents = matrix_result.get("documents", [])
    topic = matrix_result.get("topic", "")
    doc_names = [d["name"] for d in documents]

    doc = Document()

    # Title
    title = doc.add_heading(f"Matrix Compare: {topic}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Documents: {', '.join(doc_names)}")

    # Build table
    col_count = 1 + len(doc_names)
    table = doc.add_table(rows=1 + len(aspects), cols=col_count)
    table.style = "Light Grid Accent 1"

    # Header row
    hdr = table.rows[0]
    hdr.cells[0].text = "Aspect"
    for i, name in enumerate(doc_names):
        hdr.cells[i + 1].text = name

    # Data rows
    for row_idx, aspect in enumerate(aspects, start=1):
        row = table.rows[row_idx]
        row.cells[0].text = aspect
        for col_idx, doc_info in enumerate(documents, start=1):
            row.cells[col_idx].text = doc_info.get("values", {}).get(aspect, "")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Generated by Contract Dashboard")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
