"""Tests for export service."""

import pytest
from io import BytesIO

from app.models import SearchResult, QAResponse, Citation
from app.services.export import (
    export_search_results_html,
    export_search_results_docx,
    export_compare_results_html,
    export_compare_results_docx,
    export_qa_response_html,
    export_qa_response_docx,
)


# =============================================================================
# Fixtures
# =============================================================================

@pytest.fixture
def sample_search_results():
    """Create sample search results for testing."""
    return [
        SearchResult(
            file_id=1,
            file_path="/path/to/doc1.pdf",
            filename="Contract_A.pdf",
            page_number=5,
            snippet="The <mark>overtime</mark> rate shall be 1.5 times the regular rate...",
            score=12.5,
        ),
        SearchResult(
            file_id=2,
            file_path="/path/to/doc2.pdf",
            filename="Contract_B.pdf",
            page_number=10,
            snippet="All <mark>overtime</mark> work must be approved in advance...",
            score=10.2,
        ),
    ]


@pytest.fixture
def sample_compare_multi_result():
    """Create sample multi-document comparison result."""
    return {
        "documents": [
            {"file_id": 1, "filename": "Contract_A.pdf", "page_count": 50},
            {"file_id": 2, "filename": "Contract_B.pdf", "page_count": 45},
        ],
        "matches": [
            {
                "file_id": 1,
                "filename": "Contract_A.pdf",
                "page_number": 5,
                "snippet": "...the <mark>vacation</mark> policy states...",
            },
            {
                "file_id": 1,
                "filename": "Contract_A.pdf",
                "page_number": 8,
                "snippet": "...additional <mark>vacation</mark> days for...",
            },
            {
                "file_id": 2,
                "filename": "Contract_B.pdf",
                "page_number": 12,
                "snippet": "...<mark>vacation</mark> entitlement shall be...",
            },
        ],
        "topic": "vacation",
    }


@pytest.fixture
def sample_compare_legacy_result():
    """Create sample legacy two-document comparison result."""
    return {
        "doc_a": {"file_id": 1, "filename": "Contract_A.pdf"},
        "doc_b": {"file_id": 2, "filename": "Contract_B.pdf"},
        "matches_a": [
            {"page": 5, "context": "...the <mark>grievance</mark> procedure..."},
        ],
        "matches_b": [
            {"page": 10, "context": "...filing a <mark>grievance</mark> must..."},
            {"page": 11, "context": "...<mark>grievance</mark> resolution steps..."},
        ],
    }


@pytest.fixture
def sample_qa_response():
    """Create sample QA response for testing."""
    return QAResponse(
        answer="Based on the documents, the overtime rate is 1.5 times the regular hourly rate [Source 1]. Overtime must be pre-approved by management [Source 2].",
        citations=[
            Citation(
                file_id=1,
                file_path="/path/to/doc1.pdf",
                filename="Contract_A.pdf",
                page_number=5,
                cited_text="The overtime rate shall be 1.5 times the regular hourly rate for all hours worked beyond 40 per week.",
            ),
            Citation(
                file_id=2,
                file_path="/path/to/doc2.pdf",
                filename="Contract_B.pdf",
                page_number=10,
                cited_text="All overtime work must be approved in advance by the employee's direct supervisor.",
            ),
        ],
        no_evidence=False,
    )


@pytest.fixture
def sample_qa_response_no_evidence():
    """Create sample QA response with no evidence."""
    return QAResponse(
        answer="Not found in the documents provided.",
        citations=[],
        no_evidence=True,
    )


# =============================================================================
# HTML Export Tests
# =============================================================================

class TestSearchResultsHTML:
    """Tests for search results HTML export."""

    def test_export_contains_query(self, sample_search_results):
        """Test that exported HTML contains the search query."""
        html = export_search_results_html(sample_search_results, "overtime")
        assert "overtime" in html
        assert "Search Results" in html

    def test_export_contains_results(self, sample_search_results):
        """Test that exported HTML contains result information."""
        html = export_search_results_html(sample_search_results, "overtime")
        assert "Contract_A.pdf" in html
        assert "Contract_B.pdf" in html
        assert "Page 5" in html
        assert "Page 10" in html

    def test_export_preserves_highlight_marks(self, sample_search_results):
        """Test that <mark> tags are preserved in HTML output."""
        html = export_search_results_html(sample_search_results, "overtime")
        assert "<mark>overtime</mark>" in html

    def test_export_empty_results(self):
        """Test HTML export with no results."""
        html = export_search_results_html([], "nonexistent")
        assert "No results found" in html
        assert "nonexistent" in html

    def test_export_is_valid_html(self, sample_search_results):
        """Test that output is valid HTML document."""
        html = export_search_results_html(sample_search_results, "test")
        assert html.startswith("<!DOCTYPE html>")
        assert "<html" in html
        assert "</html>" in html
        assert "<head>" in html
        assert "<body>" in html


class TestCompareResultsHTML:
    """Tests for compare results HTML export."""

    def test_export_multi_document_contains_topic(self, sample_compare_multi_result):
        """Test that multi-doc comparison HTML contains topic."""
        html = export_compare_results_html(sample_compare_multi_result, "vacation")
        assert "vacation" in html
        assert "Document Comparison" in html

    def test_export_multi_document_contains_documents(self, sample_compare_multi_result):
        """Test that multi-doc comparison lists all documents."""
        html = export_compare_results_html(sample_compare_multi_result, "vacation")
        assert "Contract_A.pdf" in html
        assert "Contract_B.pdf" in html

    def test_export_multi_document_contains_matches(self, sample_compare_multi_result):
        """Test that multi-doc comparison contains match snippets."""
        html = export_compare_results_html(sample_compare_multi_result, "vacation")
        assert "Page 5" in html
        assert "Page 8" in html
        assert "Page 12" in html

    def test_export_legacy_format(self, sample_compare_legacy_result):
        """Test legacy two-document comparison format."""
        html = export_compare_results_html(sample_compare_legacy_result, "grievance")
        assert "Contract_A.pdf" in html
        assert "Contract_B.pdf" in html
        assert "grievance" in html

    def test_export_empty_comparison(self):
        """Test HTML export with empty comparison."""
        html = export_compare_results_html({}, "")
        assert "Document Comparison" in html


class TestQAResponseHTML:
    """Tests for Q&A response HTML export."""

    def test_export_contains_question(self, sample_qa_response):
        """Test that HTML contains the original question."""
        html = export_qa_response_html(sample_qa_response, "What is the overtime rate?")
        assert "What is the overtime rate?" in html

    def test_export_contains_answer(self, sample_qa_response):
        """Test that HTML contains the answer."""
        html = export_qa_response_html(sample_qa_response, "What is the overtime rate?")
        assert "1.5 times" in html

    def test_export_contains_citations(self, sample_qa_response):
        """Test that HTML contains citations."""
        html = export_qa_response_html(sample_qa_response, "test")
        assert "Contract_A.pdf" in html
        assert "Page 5" in html
        assert "Contract_B.pdf" in html
        assert "Page 10" in html

    def test_export_no_evidence_warning(self, sample_qa_response_no_evidence):
        """Test that no-evidence response shows warning."""
        html = export_qa_response_html(sample_qa_response_no_evidence, "test")
        assert "Not found in documents" in html


# =============================================================================
# DOCX Export Tests
# =============================================================================

class TestSearchResultsDOCX:
    """Tests for search results DOCX export."""

    def test_export_returns_bytes(self, sample_search_results):
        """Test that DOCX export returns bytes."""
        result = export_search_results_docx(sample_search_results, "overtime")
        assert isinstance(result, bytes)

    def test_export_is_valid_docx(self, sample_search_results):
        """Test that output is valid DOCX (ZIP-based format)."""
        result = export_search_results_docx(sample_search_results, "overtime")
        # DOCX files are ZIP archives starting with PK signature
        assert result[:2] == b'PK'

    def test_export_can_be_opened(self, sample_search_results):
        """Test that exported DOCX can be opened by python-docx."""
        from docx import Document
        result = export_search_results_docx(sample_search_results, "overtime")
        doc = Document(BytesIO(result))
        # Should have content
        assert len(doc.paragraphs) > 0

    def test_export_empty_results(self):
        """Test DOCX export with no results."""
        result = export_search_results_docx([], "nonexistent")
        assert isinstance(result, bytes)
        assert result[:2] == b'PK'


class TestCompareResultsDOCX:
    """Tests for compare results DOCX export."""

    def test_export_multi_returns_bytes(self, sample_compare_multi_result):
        """Test that multi-doc DOCX export returns bytes."""
        result = export_compare_results_docx(sample_compare_multi_result, "vacation")
        assert isinstance(result, bytes)

    def test_export_multi_is_valid_docx(self, sample_compare_multi_result):
        """Test that multi-doc output is valid DOCX."""
        result = export_compare_results_docx(sample_compare_multi_result, "vacation")
        assert result[:2] == b'PK'

    def test_export_legacy_returns_bytes(self, sample_compare_legacy_result):
        """Test that legacy format DOCX export returns bytes."""
        result = export_compare_results_docx(sample_compare_legacy_result, "grievance")
        assert isinstance(result, bytes)

    def test_export_can_be_opened(self, sample_compare_multi_result):
        """Test that exported DOCX can be opened."""
        from docx import Document
        result = export_compare_results_docx(sample_compare_multi_result, "vacation")
        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) > 0


class TestQAResponseDOCX:
    """Tests for Q&A response DOCX export."""

    def test_export_returns_bytes(self, sample_qa_response):
        """Test that DOCX export returns bytes."""
        result = export_qa_response_docx(sample_qa_response, "What is the overtime rate?")
        assert isinstance(result, bytes)

    def test_export_is_valid_docx(self, sample_qa_response):
        """Test that output is valid DOCX."""
        result = export_qa_response_docx(sample_qa_response, "test")
        assert result[:2] == b'PK'

    def test_export_can_be_opened(self, sample_qa_response):
        """Test that exported DOCX can be opened."""
        from docx import Document
        result = export_qa_response_docx(sample_qa_response, "test")
        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_export_no_evidence(self, sample_qa_response_no_evidence):
        """Test DOCX export with no evidence response."""
        result = export_qa_response_docx(sample_qa_response_no_evidence, "test")
        assert isinstance(result, bytes)
        assert result[:2] == b'PK'


# =============================================================================
# Edge Case Tests
# =============================================================================

class TestEdgeCases:
    """Test edge cases and special characters."""

    def test_html_escapes_special_chars(self):
        """Test that special HTML characters are escaped."""
        results = [
            SearchResult(
                file_id=1,
                file_path="/path/to/doc.pdf",
                filename="Test<script>alert('xss')</script>.pdf",
                page_number=1,
                snippet="Test & <script>alert('xss')</script>",
                score=1.0,
            )
        ]
        html = export_search_results_html(results, "<script>alert('xss')</script>")
        assert "<script>" not in html.replace("<mark>", "").replace("</mark>", "")
        assert "&lt;script&gt;" in html

    def test_long_query_handling(self, sample_search_results):
        """Test handling of very long queries."""
        long_query = "a" * 1000
        html = export_search_results_html(sample_search_results, long_query)
        assert isinstance(html, str)

    def test_unicode_handling(self):
        """Test handling of unicode characters."""
        results = [
            SearchResult(
                file_id=1,
                file_path="/path/to/doc.pdf",
                filename="Contrat_Francais.pdf",
                page_number=1,
                snippet="L'employe a droit a 20 jours de conge annuel...",
                score=1.0,
            )
        ]
        html = export_search_results_html(results, "conge")
        assert "Contrat_Francais.pdf" in html
        # Apostrophe is HTML-escaped in the output
        assert "L&#x27;employe" in html or "L'employe" in html
